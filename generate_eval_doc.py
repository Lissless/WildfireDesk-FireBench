import json
import csv
import re
import sys
import argparse
from collections import Counter, defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if hasattr(sys.stdout, 'reconfigure'):
	sys.stdout.reconfigure(encoding='utf-8')  # type: ignore[union-attr]


def set_cell_bg(cell, hex_color):
	tc = cell._tc
	tcPr = tc.get_or_add_tcPr()
	shd = OxmlElement("w:shd")
	shd.set(qn("w:val"), "clear")
	shd.set(qn("w:color"), "auto")
	shd.set(qn("w:fill"), hex_color)
	tcPr.append(shd)


def add_inline_bold(para, text):
	parts = re.split(r"\*\*(.+?)\*\*", text)
	for i, part in enumerate(parts):
		run = para.add_run(part)
		if i % 2 == 1:
			run.bold = True


def add_markdown_content(doc, text, base_font_size=11):
	lines = text.split("\n")
	i = 0
	while i < len(lines):
		line = lines[i]

		if line.startswith("### "):
			p = doc.add_paragraph()
			run = p.add_run(line[4:].strip())
			run.bold = True
			run.font.size = Pt(base_font_size)
			run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

		elif line.startswith("## "):
			p = doc.add_paragraph()
			run = p.add_run(line[3:].strip())
			run.bold = True
			run.font.size = Pt(base_font_size + 1)
			run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

		elif re.match(r"^\s{3,}- ", line):
			p = doc.add_paragraph(style="List Bullet 2")
			content = re.sub(r"^\s+-\s+", "", line)
			add_inline_bold(p, content)
			for run in p.runs:
				run.font.size = Pt(base_font_size)

		elif line.startswith("- "):
			p = doc.add_paragraph(style="List Bullet")
			add_inline_bold(p, line[2:])
			for run in p.runs:
				run.font.size = Pt(base_font_size)

		elif re.match(r"^\d+\.\s+\*\*", line):
			p = doc.add_paragraph(style="List Number")
			add_inline_bold(p, re.sub(r"^\d+\.\s+", "", line))
			for run in p.runs:
				run.font.size = Pt(base_font_size)

		elif line.strip() == "":
			pass

		else:
			p = doc.add_paragraph()
			add_inline_bold(p, line)
			for run in p.runs:
				run.font.size = Pt(base_font_size)

		i += 1


def add_label_paragraph(doc, label, font_size=10, color="555555"):
	p = doc.add_paragraph()
	run = p.add_run(label)
	run.bold = True
	run.font.size = Pt(font_size)
	r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
	run.font.color.rgb = RGBColor(r, g, b)
	return p


def add_horizontal_rule(doc):
	p = doc.add_paragraph()
	pPr = p._p.get_or_add_pPr()
	pBdr = OxmlElement("w:pBdr")
	bottom = OxmlElement("w:bottom")
	bottom.set(qn("w:val"), "single")
	bottom.set(qn("w:sz"), "6")
	bottom.set(qn("w:space"), "1")
	bottom.set(qn("w:color"), "CCCCCC")
	pBdr.append(bottom)
	pPr.append(pBdr)
	return p


def build_rubric_data(csv_path):
	"""从CSV读取rubric数据，返回 Category -> [行字典列表] 的映射，自动处理重复Label。"""
	category_rows = {}

	with open(csv_path, encoding='utf-8') as f:
		reader = csv.DictReader(f)
		for row in reader:
			cat = row['Category'].strip()
			if not cat:
				continue
			if cat not in category_rows:
				category_rows[cat] = []
			category_rows[cat].append({
				'Label': row['Label'].strip(),
				'Element': row['Element (Question)'].strip(),
				'Description': row['Description'].strip(),
				'Why it matters': row['Why it matters'].strip(),
				'Points': row['Points'].strip(),
			})

	# 检测重复Label并添加数字后缀，与civic_judge.json的key保持一致
	for cat, rows in category_rows.items():
		labels = [r['Label'] for r in rows]
		counts = Counter(labels)
		seen = defaultdict(int)
		for row in rows:
			label = row['Label']
			if counts[label] > 1:
				seen[label] += 1
				row['Label'] = f"{label}{seen[label]}"

	return category_rows


def add_rubric_table(doc, rubric_rows):
	"""添加6列rubric评分表格：Label, Element, Description, Why it matters, Max Points, Your Points。"""
	col_headers = ['Label', 'Element', 'Description', 'Why it matters', 'Max Points', 'Your Points']
	# 总可用宽度约6.1英寸（8.5 - 1.2*2）
	col_widths = [Inches(0.9), Inches(1.4), Inches(1.6), Inches(1.2), Inches(0.5), Inches(0.5)]

	table = doc.add_table(rows=1, cols=6)
	table.style = "Table Grid"

	# 表头行
	hdr = table.rows[0].cells
	for i, header in enumerate(col_headers):
		hdr[i].text = header
		for para in hdr[i].paragraphs:
			for run in para.runs:
				run.bold = True
				run.font.size = Pt(9)
		set_cell_bg(hdr[i], "D9E1F2")

	# 数据行
	for row_data in rubric_rows:
		row = table.add_row().cells
		row[0].text = row_data['Label']
		row[1].text = row_data['Element']
		row[2].text = row_data['Description']
		row[3].text = row_data['Why it matters']
		row[4].text = row_data['Points']
		row[5].text = ''
		for i in range(6):
			for para in row[i].paragraphs:
				for run in para.runs:
					run.font.size = Pt(9)

	# 设置列宽
	for row in table.rows:
		for i, width in enumerate(col_widths):
			row.cells[i].width = width


def generate_doc(json_path="eval/data/civic_judge.json",
				 csv_path="eval/data/civicbench_rubrics.xlsx - Rubric Questions Full.csv",
				 output_path="eval/data/civic_expert_eval.docx"):
	with open(json_path, encoding="utf-8") as f:
		data = json.load(f)

	rubric_data = build_rubric_data(csv_path)

	doc = Document()
	for section in doc.sections:
		section.top_margin = Inches(1)
		section.bottom_margin = Inches(1)
		section.left_margin = Inches(1.2)
		section.right_margin = Inches(1.2)
	title = doc.add_heading("Civic Chatbot Expert Evaluation", level=0)
	title.alignment = WD_ALIGN_PARAGRAPH.CENTER
	sub = doc.add_paragraph("Please grade conversation in each dimension for each catogory")
	sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
	sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
	doc.add_paragraph()

	for idx, item in enumerate(data):
		if idx > 0:
			doc.add_page_break()

		_id = item.get("_id", str(idx + 1))
		difficulty = item.get("difficulty", "")
		high_class = item.get("high_class", "")
		sub_class = item.get("sub_class", "")
		prompts = item["conversation"]["prompts"]
		responses = item["conversation"]["responses"]

		h = doc.add_heading(f"Question #{_id}", level=2)
		h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

		meta = doc.add_paragraph()
		meta.add_run(f"Difficulty: ").bold = True
		meta.add_run(difficulty + "    ")
		meta.add_run(f"Category: ").bold = True
		meta.add_run(high_class + "    ")
		meta.add_run(f"Subcategory: ").bold = True
		meta.add_run(sub_class)
		for run in meta.runs:
			run.font.size = Pt(9.5)

		add_horizontal_rule(doc)
		doc.add_paragraph()

		for turn_idx, (prompt, response) in enumerate(zip(prompts, responses)):
			if len(prompts) > 1:
				add_label_paragraph(doc, f"【Turn {turn_idx + 1}】", font_size=9, color="888888")

			add_label_paragraph(doc, "Question", font_size=10, color="1F497D")
			q_table = doc.add_table(rows=1, cols=1)
			q_table.style = "Table Grid"
			q_cell = q_table.rows[0].cells[0]
			set_cell_bg(q_cell, "EEF3FB")
			q_para = q_cell.paragraphs[0]
			q_para.add_run(prompt)
			q_para.runs[0].font.size = Pt(10.5)

			doc.add_paragraph()

			add_label_paragraph(doc, "Chatbot Response", font_size=10, color="375623")
			add_markdown_content(doc, response, base_font_size=10.5)
			doc.add_paragraph()

		add_horizontal_rule(doc)
		add_label_paragraph(doc, "Expert Grading", font_size=11, color="7B2C2C")
		doc.add_paragraph()

		# 根据high_class查找对应rubric行，找不到时打印警告
		rubric_rows = rubric_data.get(high_class)
		if rubric_rows is None:
			print(f"警告：未找到Category '{high_class}' 对应的rubric（id={_id}）")
			rubric_rows = []
		add_rubric_table(doc, rubric_rows)

	doc.save(output_path)
	print(f"Done：{output_path}（{len(data)} elements）")


if __name__ == "__main__":
	parser = argparse.ArgumentParser(description="生成civic评测Word文档")
	parser.add_argument("--input", default="eval/data/civic_judge.json", help="JSON路径")
	parser.add_argument("--rubric", default="eval/data/civicbench_rubrics.xlsx - Rubric Questions Full.csv", help="Rubric CSV路径")
	parser.add_argument("--output", default="eval/data/civic_expert_eval.docx", help="Word输出路径")
	args = parser.parse_args()
	generate_doc(json_path=args.input, csv_path=args.rubric, output_path=args.output)
